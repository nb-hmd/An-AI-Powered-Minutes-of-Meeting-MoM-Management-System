"""
Export Service - Generate PDF and DOCX exports of Minutes of Meeting.
"""

import os
from io import BytesIO
from datetime import datetime
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config.settings import EXPORT_DIR_PDF, EXPORT_DIR_DOCX
from models.activity_log import ActivityLogModel


class ExportService:
    """Service for generating PDF and DOCX exports of MoMs."""

    # ========================================================
    # PDF EXPORT
    # ========================================================
    @staticmethod
    def generate_pdf(mom_data: dict, user_id=None) -> BytesIO:
        """Generate a PDF for a single MoM. Returns a BytesIO buffer."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=1.5*cm, bottomMargin=1.5*cm,
            leftMargin=2*cm, rightMargin=2*cm,
        )

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "MoMTitle", parent=styles["Title"],
            fontSize=18, textColor=HexColor("#2d3748"),
            spaceAfter=6,
        )
        heading_style = ParagraphStyle(
            "MoMHeading", parent=styles["Heading2"],
            fontSize=13, textColor=HexColor("#667eea"),
            spaceBefore=12, spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "MoMBody", parent=styles["Normal"],
            fontSize=10, leading=14,
        )
        meta_style = ParagraphStyle(
            "MoMMeta", parent=styles["Normal"],
            fontSize=9, textColor=HexColor("#718096"),
        )

        story = []

        # Header
        story.append(Paragraph("MINUTES OF MEETING", title_style))
        story.append(HRFlowable(
            width="100%", thickness=2,
            color=HexColor("#667eea"), spaceAfter=12,
        ))

        # Metadata table
        meta_data = [
            ["Title:", mom_data.get("title", "N/A")],
            ["Date & Time:", str(mom_data.get("date_time", "N/A"))],
            ["Venue:", mom_data.get("venue", "N/A")],
            ["Category:", mom_data.get("category", "N/A")],
            ["Department:", mom_data.get("department", "N/A")],
            ["Created By:", mom_data.get("creator_name", "N/A")],
        ]
        meta_table = Table(meta_data, colWidths=[3*cm, 13*cm])
        meta_table.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("TEXTCOLOR", (0, 0), (0, -1), HexColor("#4a5568")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 12))

        # Agenda
        if mom_data.get("agenda"):
            story.append(Paragraph("AGENDA", heading_style))
            for line in mom_data["agenda"].split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 8))

        # Discussion
        if mom_data.get("discussion"):
            story.append(Paragraph("DISCUSSION", heading_style))
            for line in mom_data["discussion"].split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 8))

        # Decisions
        if mom_data.get("decisions"):
            story.append(Paragraph("DECISIONS", heading_style))
            for line in mom_data["decisions"].split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 8))

        # Attendees table
        attendees = mom_data.get("attendees", [])
        if attendees:
            story.append(Paragraph("ATTENDEES", heading_style))
            att_headers = ["Name", "Role", "Email", "Department"]
            att_data = [att_headers]
            for a in attendees:
                att_data.append([
                    a.get("name", ""), a.get("role", ""),
                    a.get("email", ""), a.get("department", ""),
                ])
            att_table = Table(att_data, colWidths=[4*cm, 3*cm, 5*cm, 4*cm])
            att_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#667eea")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#f7fafc")]),
            ]))
            story.append(att_table)
            story.append(Spacer(1, 8))

        # Action Items table
        action_items = mom_data.get("action_items", [])
        if action_items:
            story.append(Paragraph("ACTION ITEMS", heading_style))
            ai_headers = ["Description", "Assigned To", "Deadline", "Status"]
            ai_data = [ai_headers]
            for ai in action_items:
                ai_data.append([
                    ai.get("description", ""),
                    ai.get("assigned_to", ""),
                    str(ai.get("deadline", "")),
                    ai.get("status", "pending").replace("_", " ").title(),
                ])
            ai_table = Table(ai_data, colWidths=[6*cm, 3*cm, 3.5*cm, 3.5*cm])
            ai_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#48bb78")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#f0fff4")]),
            ]))
            story.append(ai_table)

        # Footer
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
        story.append(Paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} | Minutes of Meeting System",
            meta_style
        ))

        doc.build(story)
        buffer.seek(0)

        # Log activity
        if user_id:
            ActivityLogModel.log_activity(
                user_id=user_id, username="",
                action="MoM Exported (PDF)",
                details=f"Exported MoM: {mom_data.get('title', 'N/A')}",
            )

        return buffer

    # ========================================================
    # DOCX EXPORT
    # ========================================================
    @staticmethod
    def generate_docx(mom_data: dict, user_id=None) -> BytesIO:
        """Generate a DOCX for a single MoM. Returns a BytesIO buffer."""
        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # Title
        title_para = doc.add_heading("MINUTES OF MEETING", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

        doc.add_paragraph("")  # spacer

        # Metadata table
        meta_table = doc.add_table(rows=6, cols=2, style="Light Shading Accent 1")
        meta_items = [
            ("Title", mom_data.get("title", "N/A")),
            ("Date & Time", str(mom_data.get("date_time", "N/A"))),
            ("Venue", mom_data.get("venue", "N/A")),
            ("Category", mom_data.get("category", "N/A")),
            ("Department", mom_data.get("department", "N/A")),
            ("Created By", mom_data.get("creator_name", "N/A")),
        ]
        for i, (label, value) in enumerate(meta_items):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[1].text = str(value)
            for cell in meta_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(9)

        doc.add_paragraph("")

        # Content sections
        sections = [
            ("Agenda", mom_data.get("agenda", "")),
            ("Discussion", mom_data.get("discussion", "")),
            ("Decisions", mom_data.get("decisions", "")),
        ]
        for section_title, content in sections:
            if content:
                heading = doc.add_heading(section_title.upper(), level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
                doc.add_paragraph(content)

        # Attendees table
        attendees = mom_data.get("attendees", [])
        if attendees:
            doc.add_heading("ATTENDEES", level=2)
            att_table = doc.add_table(rows=1, cols=4, style="Medium Shading 1 Accent 1")
            att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = att_table.rows[0].cells
            for i, text in enumerate(["Name", "Role", "Email", "Department"]):
                headers[i].text = text

            for a in attendees:
                row = att_table.add_row().cells
                row[0].text = a.get("name", "")
                row[1].text = a.get("role", "")
                row[2].text = a.get("email", "")
                row[3].text = a.get("department", "")

        # Action Items table
        action_items = mom_data.get("action_items", [])
        if action_items:
            doc.add_heading("ACTION ITEMS", level=2)
            ai_table = doc.add_table(rows=1, cols=4, style="Medium Shading 1 Accent 1")
            ai_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ai_table.rows[0].cells
            for i, text in enumerate(["Description", "Assigned To", "Deadline", "Status"]):
                headers[i].text = text

            for ai in action_items:
                row = ai_table.add_row().cells
                row[0].text = ai.get("description", "")
                row[1].text = ai.get("assigned_to", "")
                row[2].text = str(ai.get("deadline", ""))
                row[3].text = ai.get("status", "pending").replace("_", " ").title()

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} | Minutes of Meeting System"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xa0, 0xae, 0xc0)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Log activity
        if user_id:
            ActivityLogModel.log_activity(
                user_id=user_id, username="",
                action="MoM Exported (DOCX)",
                details=f"Exported MoM: {mom_data.get('title', 'N/A')}",
            )

        return buffer

    # ========================================================
    # BULK EXPORT (ZIP)
    # ========================================================
    @staticmethod
    def generate_bulk_zip(moms_data_list: list, user_id=None) -> BytesIO:
        """Generate a ZIP file containing PDFs and DOCX for all provided MoMs."""
        import zipfile
        import json

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            metadata = []
            for mom in moms_data_list:
                title_safe = "".join(c for c in mom.get("title", "mom") if c.isalnum() or c in " _-")[:50]

                # PDF
                pdf_buf = ExportService.generate_pdf(mom)
                zf.writestr(f"pdf/{title_safe}.pdf", pdf_buf.read())

                # DOCX
                docx_buf = ExportService.generate_docx(mom)
                zf.writestr(f"docx/{title_safe}.docx", docx_buf.read())

                metadata.append({
                    "id": mom.get("id"),
                    "title": mom.get("title"),
                    "date_time": str(mom.get("date_time")),
                    "venue": mom.get("venue"),
                    "category": mom.get("category"),
                })

            # JSON metadata
            zf.writestr("metadata.json", json.dumps(metadata, indent=2, default=str))

        zip_buffer.seek(0)

        if user_id:
            ActivityLogModel.log_activity(
                user_id=user_id, username="",
                action="Bulk Export",
                details=f"Exported {len(moms_data_list)} MoMs as ZIP",
            )

        return zip_buffer
